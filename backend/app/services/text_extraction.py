"""
Text Extraction Service
Handles extraction of text from various resume file formats
"""

import pdfplumber
import docx
import pypandoc
import os
from typing import Optional


class TextExtractionService:
    """
    Service for extracting text from resume files.
    Supports PDF, DOCX, and RTF formats.
    """
    
    def extract_text(self, filepath: str) -> str:
        """
        Extract text from a file based on its extension.
        
        Args:
            filepath: Path to the file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is unsupported or extraction fails
        """
        file_ext = self._get_file_extension(filepath)
        
        if file_ext == '.pdf':
            return self._extract_from_pdf(filepath)
        elif file_ext == '.docx':
            return self._extract_from_docx(filepath)
        elif file_ext == '.rtf':
            return self._extract_from_rtf(filepath)
        else:
            raise ValueError(
                f"Unsupported file format: {file_ext}. "
                "Supported formats: PDF (.pdf), Word (.docx), RTF (.rtf)"
            )
    
    def _get_file_extension(self, filepath: str) -> str:
        """Get the lowercase file extension."""
        _, ext = os.path.splitext(filepath)
        return ext.lower()
    
    def _extract_from_pdf(self, pdf_path: str) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Extracted text
            
        Raises:
            ValueError: If extraction fails or no text found
        """
        try:
            text_parts = []
            
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            if not text_parts:
                raise ValueError(
                    "No text could be extracted from the PDF. "
                    "The file may be scanned or image-based. "
                    "Try using an OCR tool or converting to a text-based PDF."
                )
            
            full_text = "\n\n".join(text_parts)
            return full_text.strip()
            
        except Exception as e:
            raise ValueError(f"Error extracting text from PDF: {str(e)}")
    
    def _extract_from_docx(self, docx_path: str) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Extracted text
            
        Raises:
            ValueError: If extraction fails or no text found
        """
        try:
            doc = docx.Document(docx_path)
            
            # Extract text from paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        table_text.append(row_text)
            
            # Combine all text
            all_text = paragraphs + table_text
            
            if not all_text:
                raise ValueError(
                    "No text could be extracted from the DOCX file. "
                    "The file may be empty or corrupted."
                )
            
            full_text = "\n".join(all_text)
            return full_text.strip()
            
        except Exception as e:
            raise ValueError(f"Error extracting text from DOCX: {str(e)}")
    
    def _extract_from_rtf(self, rtf_path: str) -> str:
        """
        Extract text from an RTF file using pypandoc.
        
        Args:
            rtf_path: Path to RTF file
            
        Returns:
            Extracted text
            
        Raises:
            ValueError: If extraction fails or pandoc is not installed
        """
        try:
            # Convert RTF to plain text using pypandoc
            text = pypandoc.convert_file(rtf_path, 'plain')
            
            if not text or not text.strip():
                raise ValueError(
                    "No text could be extracted from the RTF file. "
                    "The file may be empty or corrupted."
                )
            
            return text.strip()
            
        except OSError as e:
            raise ValueError(
                "Pandoc is not installed on the system. "
                "RTF file extraction requires Pandoc. "
                "Please install Pandoc or convert your resume to PDF or DOCX format."
            )
        except Exception as e:
            raise ValueError(f"Error extracting text from RTF: {str(e)}")
    
    def validate_file(self, filepath: str, max_size_mb: int = 10) -> bool:
        """
        Validate that a file exists and is within size limits.
        
        Args:
            filepath: Path to the file
            max_size_mb: Maximum file size in megabytes
            
        Returns:
            True if valid
            
        Raises:
            ValueError: If validation fails
        """
        if not os.path.exists(filepath):
            raise ValueError(f"File not found: {filepath}")
        
        file_size_mb = os.path.getsize(filepath) / (1024 * 1024)
        if file_size_mb > max_size_mb:
            raise ValueError(
                f"File size ({file_size_mb:.2f}MB) exceeds maximum allowed size ({max_size_mb}MB)"
            )
        
        return True
    
    def get_file_info(self, filepath: str) -> dict:
        """
        Get information about a file.
        
        Args:
            filepath: Path to the file
            
        Returns:
            Dictionary with file information
        """
        if not os.path.exists(filepath):
            return {"error": "File not found"}
        
        file_stats = os.stat(filepath)
        _, ext = os.path.splitext(filepath)
        
        return {
            "filename": os.path.basename(filepath),
            "extension": ext.lower(),
            "size_bytes": file_stats.st_size,
            "size_mb": round(file_stats.st_size / (1024 * 1024), 2),
            "exists": True
        }